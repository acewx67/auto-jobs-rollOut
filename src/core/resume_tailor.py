"""
Core resume tailoring pipeline - Main orchestration layer.

Coordinates all components:
- Resume parsing
- Job description analysis
- ATS optimization
- AI-powered tailoring
- Output generation

This is the main entry point for resume tailoring operations.
"""

import os
import json
import logging
from pathlib import Path
from typing import Dict, Optional, Tuple
from datetime import datetime

from src.core.resume_parser import ResumeParser, ResumeParseError
from src.utils.ats_optimizer import ATSOptimizer
from src.utils.latex_generator import LatexResumeGenerator, LatexGeneratorError
from src.groq_client.client import GroqClient, GroqClientError

logger = logging.getLogger(__name__)


class TailorPipelineError(Exception):
    """Custom exception for pipeline errors"""
    pass


class TailorPipeline:
    """
    Main orchestration for resume tailoring.
    
    Manages the complete workflow:
    1. Parse resume file (PDF/DOCX)
    2. Analyze job description with Groq
    3. Calculate initial ATS score
    4. Generate AI-tailored resume
    5. Apply ATS optimizations
    6. Calculate final ATS score
    7. Generate comprehensive report
    8. Export tailored resume
    
    Example usage:
        pipeline = TailorPipeline()
        result = pipeline.tailor(resume_path, job_desc, output_format="pdf")
        print(f"ATS Score: {result['final_ats_score']}")
    """
    
    def __init__(self, groq_api_key: Optional[str] = None, 
                 output_dir: str = "data/resumes/output"):
        """
        Initialize the tailoring pipeline.
        
        Args:
            groq_api_key (str, optional): Groq API key. If not provided,
                                         uses GROQ_API_KEY environment variable
            output_dir (str): Directory to save tailored resumes
        """
        try:
            self.parser = ResumeParser()
            self.optimizer = ATSOptimizer()
            self.groq_client = GroqClient(api_key=groq_api_key)
            self.generator = LatexResumeGenerator()
            
            self.output_dir = Path(output_dir)
            self.output_dir.mkdir(parents=True, exist_ok=True)
            
            self.logger = logging.getLogger(__name__)
            self.logger.info("TailorPipeline initialized successfully")
            
        except Exception as e:
            raise TailorPipelineError(f"Failed to initialize pipeline: {str(e)}")
    
    def tailor(self, resume_path: str, job_description: str,
               output_format: str = "txt") -> Dict[str, any]:
        """
        Main entry point: Tailor a resume for a specific job.
        
        Args:
            resume_path (str): Path to resume file (PDF or DOCX)
            job_description (str): Job posting text or path to job file
            output_format (str): Output format ("txt", "json", "pdf", "docx")
                                Currently supports "txt" and "json"
            
        Returns:
            Dict containing:
                - 'success': Boolean indicating if process succeeded
                - 'original_ats_score': ATS score before tailoring
                - 'tailored_resume': Full tailored resume text
                - 'final_ats_score': ATS score after tailoring
                - 'ats_improvement': Percentage improvement in score
                - 'key_changes': List of modifications made
                - 'matched_keywords': Keywords from job found in resume
                - 'missing_keywords': Important keywords still missing
                - 'metrics_preserved': Important metrics found in resume
                - 'output_file': Path to saved output file
                - 'report': Comprehensive optimization report
                - 'timestamp': Processing timestamp
                
        Raises:
            TailorPipelineError: If any step in the pipeline fails
        """
        start_time = datetime.now()
        self.logger.info(f"Starting tailor pipeline for resume: {resume_path}")
        
        try:
            # Step 1: Parse resume file
            self.logger.info("Step 1/6: Parsing resume...")
            parsed_resume = self._parse_resume(resume_path)
            
            # Step 2: Load job description (file or string)
            self.logger.info("Step 2/6: Loading job description...")
            job_desc = self._load_job_description(job_description)
            
            # Step 3: Calculate initial ATS score
            self.logger.info("Step 3/6: Calculating initial ATS score...")
            initial_ats = self.optimizer.calculate_ats_score(
                parsed_resume['normalized_text'],
                job_desc,
                parsed_resume['sections']
            )
            
            # Step 4: Analyze job with Groq
            self.logger.info("Step 4/6: Analyzing job description with AI...")
            job_analysis = self.groq_client.analyze_job_description(job_desc)
            
            # Step 5/6: Generating tailored resume with AI...
            self.logger.info("Step 5/6: Generating tailored resume with AI...")
            
            # Initial generation
            tailored_response = self.groq_client.generate_tailored_resume(
                parsed_resume['normalized_text'], 
                job_desc,
                job_analysis
            )
            
            # Extract tailored resume text
            tailored_text = self._extract_tailored_text(tailored_response)
            
            # Essential sections to check
            essential_sections = ['SUMMARY', 'EXPERIENCE', 'EDUCATION', 'PROJECTS', 'TECHNICAL SKILLS']
            
            def get_missing(text):
                missing = []
                u_text = text.upper()
                for s in essential_sections:
                    if s not in u_text:
                        missing.append(s)
                return missing
            
            missing_sections = get_missing(tailored_text)
            
            # One-time retry if sections are missing
            if missing_sections:
                self.logger.warning(f"Tailored resume is missing essential sections: {', '.join(missing_sections)}. Retrying once...")
                
                feedback = f"The previous attempt was missing the following essential sections: {', '.join(missing_sections)}. " \
                           f"Please regenerate the full resume and ensure ALL core sections are included this time."
                
                tailored_response = self.groq_client.generate_tailored_resume(
                    parsed_resume['normalized_text'], 
                    job_desc,
                    job_analysis,
                    retry_feedback=feedback
                )
                tailored_text = self._extract_tailored_text(tailored_response)
                
                # Check again after retry
                still_missing = get_missing(tailored_text)
                if still_missing:
                    self.logger.error(f"Tailored resume still missing sections after retry: {', '.join(still_missing)}")
                else:
                    self.logger.info("Retry successful: all essential sections are now present")
            
            # Apply ATS formatting
            tailored_text = self.optimizer.format_for_ats(tailored_text)
            
            # Step 6: Calculate final ATS score
            self.logger.info("Step 6/6: Calculating final ATS score...")
            final_ats = self.optimizer.calculate_ats_score(tailored_text, job_desc)
            
            # Calculate improvement
            improvement = final_ats['overall_score'] - initial_ats['overall_score']
            improvement_pct = (improvement / initial_ats['overall_score'] * 100) if initial_ats['overall_score'] > 0 else 0
            
            # Generate report
            report = self.optimizer.generate_optimization_report(
                final_ats,
                tailored_text,
                job_desc
            )
            
            # Save output
            output_file = self._save_output(tailored_text, parsed_resume, job_desc, output_format, self.output_dir)
            
            # Compile result
            result = {
                'success': True,
                'original_ats_score': initial_ats['overall_score'],
                'final_ats_score': final_ats['overall_score'],
                'ats_improvement': improvement,
                'ats_improvement_percentage': round(improvement_pct, 1),
                'tailored_resume': tailored_text,
                'key_changes': tailored_response.get('key_changes', []),
                'matched_keywords': final_ats['matched_keywords'],
                'missing_keywords': final_ats['missing_keywords'],
                'metrics_preserved': self.optimizer.preserve_metrics(tailored_text),
                'output_file': str(output_file),
                'report': report,
                'job_analysis': job_analysis,
                'timestamp': datetime.now().isoformat(),
                'processing_time_seconds': (datetime.now() - start_time).total_seconds(),
            }
            
            self.logger.info(f"Pipeline complete. ATS Score: {initial_ats['overall_score']} → {final_ats['overall_score']}")
            return result
            
        except (ResumeParseError, GroqClientError) as e:
            self.logger.error(f"Pipeline error: {str(e)}")
            raise TailorPipelineError(f"Tailoring failed: {str(e)}")
        except Exception as e:
            self.logger.error(f"Unexpected pipeline error: {str(e)}")
            raise TailorPipelineError(f"Unexpected error: {str(e)}")
    
    def _parse_resume(self, resume_path: str) -> Dict:
        """Parse resume file and return parsed content"""
        parsed = self.parser.parse(resume_path)
        self.logger.debug(f"Parsed resume: {parsed['character_count']} characters")
        return parsed
    
    def _load_job_description(self, job_input: str) -> str:
        """
        Load job description from string or file.
        
        Args:
            job_input (str): Job description text or file path
            
        Returns:
            str: Job description content
            
        Raises:
            TailorPipelineError: If file not found or empty
        """
        # Check if it's a file path
        job_path = Path(job_input)
        if job_path.exists() and job_path.is_file():
            content = job_path.read_text(encoding='utf-8')
            if not content.strip():
                raise TailorPipelineError("Job description file is empty")
            return content
        
        # Otherwise treat as raw text
        if not job_input or not job_input.strip():
            raise TailorPipelineError("Job description is empty")
        
        return job_input
    
    def _extract_tailored_text(self, response: Dict) -> str:
        """
        Extract tailored resume text from Groq response.
        
        Args:
            response (Dict): Response from generate_tailored_resume()
            
        Returns:
            str: Tailored resume text
            
        Raises:
            TailorPipelineError: If response is invalid
        """
        if 'tailored_resume' not in response:
            raise TailorPipelineError("Invalid response format from AI: missing 'tailored_resume'")
        
        text = response['tailored_resume']
        if not text or not text.strip():
            raise TailorPipelineError("AI generated empty resume")
        
        return text.strip()
    
    def _save_output(self, content: str, original_data: Dict, job_desc: str, 
                    output_format: str, result_dir: Path) -> Path:
        """
        Save tailored resume to file in specified format.
        
        Args:
            content (str): Tailored resume text
            original_data (Dict): Original parsed resume data
            job_desc (str): Job description text
            output_format (str): Desired output format
            result_dir (Path): Directory to save results
            
        Returns:
            Path: Path to the primary output file
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        source_name = Path(original_data.get('source_path', 'resume')).stem
        
        if output_format == "docx":
            output_path = result_dir / f"{source_name}_tailored_{timestamp}.docx"
            self._save_as_docx(content, output_path)
            return output_path
            
        elif output_format == "json":
            output_path = result_dir / f"{source_name}_tailored_{timestamp}.json"
            data = {
                'tailored_resume': content, 
                'timestamp': timestamp,
                'candidate_name': original_data.get('name')
            }
            output_path.write_text(json.dumps(data, indent=2), encoding='utf-8')
            return output_path
            
        elif output_format == 'latex':
            latex_code = self.generator.generate_latex(content, name=original_data.get('name'))
            output_path = result_dir / f"{source_name}_tailored_{timestamp}.tex"
            output_path.write_text(latex_code, encoding='utf-8')
            return output_path
            
        elif output_format == "latex-pdf":
            output_path_base = result_dir / f"{source_name}_tailored_{timestamp}"
            pdf_path = self.generator.generate_pdf(content, str(output_path_base), name=original_data.get('name'))
            if pdf_path:
                return Path(pdf_path)
            return output_path_base.with_suffix('.tex')
            
        else:  # Default to txt
            output_path = result_dir / f"{source_name}_tailored_{timestamp}.txt"
            output_path.write_text(content, encoding='utf-8')
            return output_path
    
    def _save_as_docx(self, content: str, output_path: Path) -> None:
        """
        Save resume content as a DOCX file with professional formatting.
        
        Args:
            content (str): Resume text
            output_path (Path): Output file path
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        except ImportError:
            self.logger.warning("python-docx not installed, saving as TXT instead")
            output_path.with_suffix('.txt').write_text(content, encoding='utf-8')
            return
        
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Parse resume into sections
        lines = content.split('\n')
        current_section = None
        
        for line in lines:
            line_stripped = line.strip()
            
            if not line_stripped:
                continue
            
            # Detect section headers (usually ALL CAPS or followed by dashes)
            if (line_stripped.isupper() and 
                len(line_stripped) < 40 and 
                not line_stripped.startswith('•')):
                # Add section heading
                heading = doc.add_heading(line_stripped, level=1)
                heading_format = heading.paragraph_format
                heading_format.space_before = Pt(6)
                heading_format.space_after = Pt(3)
                current_section = line_stripped
            
            elif line_stripped.startswith('•'):
                # Add bullet point (remove bullet and add as list)
                text = line_stripped[1:].strip()
                para = doc.add_paragraph(text, style='List Bullet')
                para.paragraph_format.space_after = Pt(2)
            
            else:
                # Regular paragraph
                if line_stripped:
                    para = doc.add_paragraph(line_stripped)
                    para.paragraph_format.space_after = Pt(2)
        
        doc.save(str(output_path))
        self.logger.info(f"DOCX document saved to: {output_path}")
    
    
    def batch_tailor(self, resume_path: str, job_files_directory: str,
                    output_format: str = "txt") -> Dict[str, any]:
        """
        Tailor a single resume against multiple job descriptions.
        
        Useful for applying to multiple positions.
        
        Args:
            resume_path (str): Path to resume file
            job_files_directory (str): Directory containing job description files
            output_format (str): Output format
            
        Returns:
            Dict with results for each job
        """
        jobs_dir = Path(job_files_directory)
        
        if not jobs_dir.exists() or not jobs_dir.is_dir():
            raise TailorPipelineError(f"Jobs directory not found: {job_files_directory}")
        
        job_files = list(jobs_dir.glob("*.txt")) + list(jobs_dir.glob("*.md"))
        
        if not job_files:
            raise TailorPipelineError(f"No job files found in {job_files_directory}")
        
        self.logger.info(f"Batch tailoring: {len(job_files)} jobs")
        
        results = {
            'total_jobs': len(job_files),
            'successful': 0,
            'failed': 0,
            'average_ats_improvement': 0,
            'tailored_resumes': {},
            'errors': {}
        }
        
        total_improvement = 0
        
        for job_file in job_files:
            try:
                self.logger.info(f"Processing: {job_file.name}")
                result = self.tailor(resume_path, str(job_file), output_format)
                
                results['tailored_resumes'][job_file.stem] = result
                results['successful'] += 1
                total_improvement += result['ats_improvement']
                
            except Exception as e:
                self.logger.error(f"Failed to tailor for {job_file.name}: {str(e)}")
                results['failed'] += 1
                results['errors'][job_file.name] = str(e)
        
        if results['successful'] > 0:
            results['average_ats_improvement'] = round(
                total_improvement / results['successful'], 1
            )
        
        self.logger.info(f"Batch complete: {results['successful']} successful, {results['failed']} failed")
        return results


class PipelineConfig:
    """Configuration for tailoring pipeline"""
    
    # Resume storage paths
    INPUT_RESUME_DIR = "data/resumes/input"
    OUTPUT_RESUME_DIR = "data/resumes/output"
    JOB_DESCRIPTIONS_DIR = "data/job_posts"
    
    # Output paths
    REPORTS_DIR = "data/reports"
    LOGS_DIR = "logs"
    
    # Pipeline settings
    DEFAULT_OUTPUT_FORMAT = "txt"
    SUPPORTED_FORMATS = ["txt", "json", "pdf", "docx", "latex-pdf"]
    
    @classmethod
    def initialize_directories(cls):
        """Create required directories if they don't exist"""
        directories = [
            cls.INPUT_RESUME_DIR,
            cls.OUTPUT_RESUME_DIR,
            cls.JOB_DESCRIPTIONS_DIR,
            cls.REPORTS_DIR,
            cls.LOGS_DIR,
        ]
        
        for directory in directories:
            Path(directory).mkdir(parents=True, exist_ok=True)
        
        logger.info("Pipeline directories initialized")

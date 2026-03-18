"""
Resume parsing module for extracting text from various resume formats.

Supports PDF and DOCX formats with automatic text normalization and
basic structural parsing into sections.
"""

import re
import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from pypdf import PdfReader
from docx import Document

logger = logging.getLogger(__name__)


class ResumeParseError(Exception):
    """Custom exception for resume parsing errors"""
    pass


class ResumeParser:
    """
    Handles parsing and extraction of resume content from various formats.
    
    Supports:
    - PDF files (.pdf)
    - DOCX files (.docx, .doc)
    
    Provides methods to:
    - Parse files in supported formats
    - Normalize and clean text
    - Extract and structure resume sections
    """
    
    # Resume section keywords for structuring
    SECTION_KEYWORDS = {
        'contact': ['contact', 'contact information', 'phone', 'email'],
        'summary': ['summary', 'professional summary', 'objective', 'profile'],
        'experience': ['experience', 'work experience', 'employment', 'professional experience'],
        'education': ['education', 'academic', 'degrees', 'certifications', 'certificate'],
        'skills': ['skills', 'technical skills', 'competencies', 'expertise'],
        'projects': ['projects', 'portfolio'],
        'languages': ['languages', 'language proficiency'],
        'certifications': ['certifications', 'licenses', 'certificates', 'awards'],
    }
    
    def __init__(self):
        """Initialize the resume parser"""
        self.logger = logging.getLogger(__name__)
    
    def parse(self, file_path: str) -> Dict[str, any]:
        """
        Parse a resume file and return extracted content.
        
        Args:
            file_path (str): Path to the resume file (PDF or DOCX)
            
        Returns:
            Dict containing:
                - 'raw_text': Full extracted text
                - 'normalized_text': Cleaned and normalized text
                - 'sections': Dict of structured resume sections
                - 'file_format': Format of the parsed file
                - 'character_count': Number of characters in resume
                
        Raises:
            ResumeParseError: If file format is unsupported or parsing fails
        """
        path = Path(file_path)
        
        if not path.exists():
            raise ResumeParseError(f"Resume file not found: {file_path}")
        
        file_ext = path.suffix.lower()
        
        try:
            if file_ext == '.pdf':
                raw_text = self._parse_pdf(file_path)
                file_format = 'pdf'
            elif file_ext in ['.docx', '.doc']:
                raw_text = self._parse_docx(file_path)
                file_format = 'docx'
            else:
                raise ResumeParseError(
                    f"Unsupported file format: {file_ext}. Supported: .pdf, .docx, .doc"
                )
            
            # Normalize the extracted text
            normalized_text = self._normalize_text(raw_text)
            
            # Structure the resume into sections
            sections = self._structure_resume(normalized_text)
            
            self.logger.info(f"Successfully parsed resume: {path.name}")
            
            return {
                'raw_text': raw_text,
                'normalized_text': normalized_text,
                'sections': sections,
                'file_format': file_format,
                'character_count': len(raw_text),
            }
            
        except ResumeParseError:
            raise
        except Exception as e:
            self.logger.error(f"Error parsing resume {file_path}: {str(e)}")
            raise ResumeParseError(f"Failed to parse resume: {str(e)}")
    
    def _parse_pdf(self, file_path: str) -> str:
        """
        Extract text from a PDF file.
        
        Args:
            file_path (str): Path to PDF file
            
        Returns:
            str: Extracted text from PDF
        """
        text = []
        try:
            reader = PdfReader(file_path)
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
                    self.logger.debug(f"Extracted page {page_num + 1} from PDF")
            
            if not text:
                raise ResumeParseError("No text found in PDF file")
                
            return "\n".join(text)
            
        except Exception as e:
            raise ResumeParseError(f"PDF parsing failed: {str(e)}")
    
    def _parse_docx(self, file_path: str) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path (str): Path to DOCX file
            
        Returns:
            str: Extracted text from DOCX
        """
        text = []
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text.append(" | ".join(row_text))
            
            if not text:
                raise ResumeParseError("No text found in DOCX file")
                
            return "\n".join(text)
            
        except Exception as e:
            raise ResumeParseError(f"DOCX parsing failed: {str(e)}")
    
    def _normalize_text(self, text: str) -> str:
        """
        Normalize and clean resume text.
        
        Operations:
        - Remove extra whitespace
        - Standardize line breaks
        - Remove control characters
        - Fix common encoding issues
        
        Args:
            text (str): Raw resume text
            
        Returns:
            str: Normalized text
        """
        # Remove control characters
        text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\t\r')
        
        # Replace multiple spaces with single space
        text = re.sub(r' +', ' ', text)
        
        # Normalize line breaks (max 2 consecutive newlines)
        text = re.sub(r'\n\n+', '\n\n', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def _structure_resume(self, text: str) -> Dict[str, str]:
        """
        Attempt to structure resume into common sections.
        
        Uses keyword matching to identify and extract resume sections.
        
        Args:
            text (str): Normalized resume text
            
        Returns:
            Dict[section_name, section_content]: Structured resume sections
        """
        sections = {}
        lines = text.split('\n')
        
        # Create a mapping of line index to content
        content_map = {}
        current_section = 'other'
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            
            # Check if this line is a section header
            matched_section = False
            for section_name, keywords in self.SECTION_KEYWORDS.items():
                if any(keyword in line_lower for keyword in keywords):
                    current_section = section_name
                    matched_section = True
                    break
            
            if current_section not in content_map:
                content_map[current_section] = []
            
            if not matched_section or line.strip():  # Add content lines
                content_map[current_section].append(line)
        
        # Clean up and return sections
        for section, lines_list in content_map.items():
            section_text = '\n'.join(lines_list).strip()
            if section_text:
                sections[section] = section_text
        
        self.logger.debug(f"Structured resume into {len(sections)} sections")
        return sections
    
    def get_text_by_section(self, parsed_resume: Dict, section: str) -> Optional[str]:
        """
        Get text from a specific section of parsed resume.
        
        Args:
            parsed_resume (Dict): Output from parse() method
            section (str): Section name
            
        Returns:
            str: Section content or None if not found
        """
        return parsed_resume.get('sections', {}).get(section)
    
    def get_all_text(self, parsed_resume: Dict) -> str:
        """
        Get full normalized text from parsed resume.
        
        Args:
            parsed_resume (Dict): Output from parse() method
            
        Returns:
            str: Full normalized resume text
        """
        return parsed_resume.get('normalized_text', '')
